"""Builds 'Documentation/Sprint Plan.docx' for DataForge.

Run from project root:  python tests/_build_sprint_doc.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Documentation" / "Sprint Plan.docx"


def shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "888888")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x2C, 0x5A, 0x99)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = bold


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.runs[0] if p.runs else p.add_run("")
    r.text = text
    r.font.size = Pt(11)


def build():
    doc = Document()

    # base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ----------------- Title -----------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("DataForge — Sprint Plan & Integration Guide")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Written in simple words. Read top to bottom.")
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ----------------- About -----------------
    h2(doc, "1. What this document is")
    para(doc,
         "This is the team plan for building DataForge in sprints. "
         "Each sprint is two weeks long. The plan shows what the team will do, "
         "how parts will join together, and how we know each part works.")

    para(doc,
         "We use easy words on purpose. Anyone on the team should be able to "
         "open this file and know what to do next.")

    # ----------------- Roles -----------------
    h2(doc, "2. Who does what")
    role_rows = [
        ("Role", "Person", "Main job"),
        ("Scrum Lead", "Atanas Todorov", "Run sprint meetings. Keep the team on plan."),
        ("Back-end Dev", "Dimitar Yanakiev", "Logic layer. Data store. File I/O."),
        ("Back-end Dev", "Georgi Dinkov", "Sort, search, recursion, validation."),
        ("QA Engineer", "Stoqn Savakov", "Write and run test cases. Open bug reports."),
        ("QA & Front-end", "Dimitar Dimitrov", "UI tests. ImGui panels. Theme."),
    ]
    add_table(doc, role_rows, widths=[3.5, 4.5, 8.5])

    # ----------------- Sprint goals table -----------------
    h2(doc, "3. Sprint plan at a glance")
    para(doc,
         "The big table below is the main view of the plan. Read each row as one "
         "sprint. The last two columns explain how the new parts will be joined "
         "to the parts that came before.")

    headers = [
        "Sprint",
        "Dates",
        "Goal",
        "What we build",
        "How we test it",
        "How we join it to the rest",
        "Done when",
    ]
    rows = [
        (
            "Sprint 0",
            "Week 1–2",
            "Set up the project.",
            "Make the empty Visual Studio solution. Add Dear ImGui as a vendored copy. Make the main window run. Pick coding style.",
            "Open the app. See a blank window. Close the app with no error.",
            "Nothing to join yet. We only set the ground.",
            "App opens and closes clean on Windows 10 and 11.",
        ),
        (
            "Sprint 1",
            "Week 3–4",
            "Make the data layer.",
            "Add Task struct. Add TaskStore. Add load and save for `.dftasks` files. Add an atomic write.",
            "Write unit tests for save and load. Use sample files from `Resources/`.",
            "Logic layer will call only these functions. No UI yet, so we test from a small console driver.",
            "We can save 1000 tasks and load them back with no loss.",
        ),
        (
            "Sprint 2",
            "Week 5–6",
            "Make the logic layer.",
            "Add validation rules. Add bubble sort and quick sort. Add linear search and recursive binary search. Add the five recursive aggregate functions.",
            "Run unit tests on each rule and each algorithm. Use both empty and full stores.",
            "Logic calls data. UI not added yet. We will wire it in Sprint 3.",
            "All logic tests pass. Bubble and quick sort give the same final order.",
        ),
        (
            "Sprint 3",
            "Week 7–8",
            "Wire the UI to the logic.",
            "Add the main task table. Add the new-task dialog. Add the edit-task dialog. Add the delete confirm dialog. Add status bar.",
            "Run the manual QA cases in `tests/task_crud.md` and `tests/hierarchy.md`.",
            "UI calls logic only. We add a small `UiState` struct as a bridge so the UI does not touch data structs by hand.",
            "A user can add, edit, and delete tasks with no crash.",
        ),
        (
            "Sprint 4",
            "Week 9–10",
            "Add search, filter, sort UI.",
            "Add the search box. Add the filter menu. Add the sort menu. Add the sort benchmark panel.",
            "Run cases in `tests/search.md`, `tests/filter.md`, `tests/sort.md`, `tests/benchmark.md`.",
            "We add a thin adapter in the presentation layer that turns user clicks into logic calls. No new logic code.",
            "Search, filter, and both sorts work on a store of 5000 tasks.",
        ),
        (
            "Sprint 5",
            "Week 11–12",
            "Add stats, theme, shortcuts, productivity dashboard.",
            "Add the stats panel. Add light and dark theme. Add all keyboard shortcuts. Add the productivity tab with pie chart and history.",
            "Run cases in `tests/statistics.md`, `tests/theme.md`, `tests/shortcuts.md`, `tests/productivity.md`.",
            "Theme and shortcuts plug into `UiState`. Productivity reads from logic aggregates only.",
            "Pie chart updates when tasks change. Theme persists across runs.",
        ),
        (
            "Sprint 6",
            "Week 13–14",
            "Polish and ship.",
            "Fix bugs from QA. Add toasts. Add help panel. Add voice input. Make the installer.",
            "Full regression run of every file in `tests/`. Smoke test on a clean Windows VM.",
            "No new layers added. Only small fixes that respect the three-tier rule.",
            "All High-priority QA cases pass. Installer runs on a fresh Windows 11 box.",
        ),
    ]
    add_table(doc, [headers, *rows], widths=[1.6, 1.6, 2.4, 4.0, 3.4, 3.4, 2.6], shrink_font=True)

    # ----------------- Integration processes -----------------
    h2(doc, "4. Example integration processes")
    para(doc,
         "An integration process is the set of steps that join two parts of the "
         "app so they work as one. Below are the most common ones for DataForge. "
         "Each one ends with a small checklist.")

    # 4.1 logic <-> data
    h3(doc, "4.1. Joining the logic layer to the data layer")
    para(doc,
         "The data layer holds the store and the file I/O. The logic layer "
         "holds the rules. Logic must call data. Data must not call logic.")
    plan_block(doc, [
        "Step 1. Write the data function (for example: `addTask`, `removeTask`).",
        "Step 2. Cover it with a unit test that uses a temp file.",
        "Step 3. Expose the function in a header inside `include/data/`.",
        "Step 4. In the logic layer, write a thin wrapper that calls the data function and adds the rule (for example: cascade delete, validation).",
        "Step 5. Add a unit test for the logic wrapper.",
        "Step 6. Run all data and logic tests. Both must pass.",
    ])
    para(doc, "Checklist:", bold=True)
    bullet(doc, "Data function has its own unit test.")
    bullet(doc, "Logic wrapper has its own unit test.")
    bullet(doc, "No `#include` from logic header inside data code.")

    # 4.2 ui <-> logic
    h3(doc, "4.2. Joining the UI to the logic layer")
    para(doc,
         "ImGui panels live in the presentation layer. They must call logic, "
         "not data. We use `UiState` as the bridge.")
    plan_block(doc, [
        "Step 1. Add the new field to `UiState` (for example: `searchQuery`).",
        "Step 2. In the panel code, read and write this field on user input.",
        "Step 3. Once per frame, call a logic function with the field as input (for example: `applySearch`).",
        "Step 4. Use the result to draw the table.",
        "Step 5. Run the manual QA case for that panel.",
    ])
    para(doc, "Checklist:", bold=True)
    bullet(doc, "Panel does not include any data header.")
    bullet(doc, "Logic call is pure: same input gives same output.")
    bullet(doc, "QA case for the panel is added under `tests/`.")

    # 4.3 persistence
    h3(doc, "4.3. Joining a new field into the save format")
    para(doc,
         "When we add a new task field (for example: tags), we must update the "
         "file format. The format is simple text so we can read it in Notepad.")
    plan_block(doc, [
        "Step 1. Add the field to the `Task` struct.",
        "Step 2. Update the serializer to write the field at the end of each task block. Old fields keep their order.",
        "Step 3. Update the parser. If the field is missing, use a default. This way old files still load.",
        "Step 4. Bump the file version line at the top of the file.",
        "Step 5. Save a sample file with the new field. Open it on a clean build to be sure.",
    ])
    para(doc, "Checklist:", bold=True)
    bullet(doc, "Old `.dftasks` files still load with no error.")
    bullet(doc, "New files load on both Debug and Release builds.")
    bullet(doc, "Atomic write still works after the format change.")

    # 4.4 algorithm
    h3(doc, "4.4. Joining a new algorithm into the sort or search menu")
    plan_block(doc, [
        "Step 1. Write the new algorithm as a free function in the logic layer.",
        "Step 2. Add unit tests for empty, one, and many tasks.",
        "Step 3. Add the option to the `SortKind` enum in the header.",
        "Step 4. In the UI sort menu, add the new entry. It calls the same dispatcher as the other sorts.",
        "Step 5. Add it to the benchmark panel so we can see how fast it runs.",
        "Step 6. Run the manual QA case in `tests/sort.md` or `tests/search.md`.",
    ])
    para(doc, "Checklist:", bold=True)
    bullet(doc, "Output matches the existing algorithm on the same input.")
    bullet(doc, "Benchmark runs without freezing the UI.")

    # 4.5 ci-like local pipeline
    h3(doc, "4.5. Joining a new commit into the main branch")
    plan_block(doc, [
        "Step 1. Pull the latest `main` branch.",
        "Step 2. Build Debug and Release with `/W4 /WX`. Fix any warning before going on.",
        "Step 3. Run all unit tests.",
        "Step 4. Run the High-priority QA cases for the area you touched.",
        "Step 5. Open a pull request. The PR body lists the QA cases that passed.",
        "Step 6. After one review, merge into `main`.",
    ])
    para(doc, "Checklist:", bold=True)
    bullet(doc, "No warnings.")
    bullet(doc, "All unit tests green.")
    bullet(doc, "At least one teammate reviewed the diff.")

    # ----------------- Risks -----------------
    h2(doc, "5. Risks and how we handle them")
    risk_rows = [
        ("Risk", "Why it can happen", "How we handle it"),
        ("Save file gets broken on crash.", "Power loss while writing.", "Use the write-and-rename trick. Test it on purpose in `PERS-03`."),
        ("Sort is too slow on big stores.", "Bubble sort is O(N²).", "Default to quick sort for big stores. Show timings in the benchmark panel."),
        ("UI freezes during a long task.", "Heavy work runs on the UI thread.", "Limit per-frame work. Use a budget of 16 ms. Move long jobs to a worker thread if needed."),
        ("Old `.dftasks` files no longer load.", "File format changes.", "Keep a version line. Parser falls back to defaults for new fields."),
        ("Theme switch leaves stale colors.", "Cached palettes inside panels.", "Re-bake palette on theme change. Covered by `THEME-03`."),
    ]
    add_table(doc, risk_rows, widths=[4.5, 5.5, 6.5])

    # ----------------- Done definition -----------------
    h2(doc, "6. Definition of done")
    bullet(doc, "Code builds clean with `/W4 /WX` in Release.")
    bullet(doc, "Unit tests for the new code pass.")
    bullet(doc, "Manual QA cases for the touched area pass.")
    bullet(doc, "No new TODO left in the code that is not linked to an issue.")
    bullet(doc, "README and this sprint plan are updated if the user-facing behavior changed.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote: {OUT}")


def add_table(doc, rows, widths, shrink_font=False):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9 if shrink_font else 10)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade(cell, "2C5A99")
            else:
                if r_idx % 2 == 0:
                    shade(cell, "F2F5FA")
            set_cell_borders(cell)


def plan_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="List Number")
        r = p.runs[0] if p.runs else p.add_run("")
        r.text = line
        r.font.size = Pt(11)


if __name__ == "__main__":
    build()
